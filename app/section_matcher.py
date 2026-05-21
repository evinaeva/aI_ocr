"""
Section matcher: parse DOCX/TXT, score candidates, pick best match.

strict_equal / soft_equal are computed via the unified `compare_lines`
primitive, so reordered lines no longer force MANUAL when the OCR text
still matches the reference character-by-character.

A TXT is treated as a structured (sectioned) document only when at
least two explicit numbered headers are present. A single header-shaped
line (e.g. `10 JETON` in a banner copy) is treated as ordinary content
and the whole TXT becomes one UNKNOWN section.
"""
import io
import re
from dataclasses import dataclass, field
from typing import List, Optional

from .normalizer import (
    normalize_strict,
    normalize_soft,
    has_placeholder,
    compare_lines,
)

# ─── Keyword scoring ────────────────────────────────────────
HIGH_PRIORITY_NAMES = {"banner", "pic", "im", "popup"}
PENALTY_NAMES = {"news", "email", "letter", "subject"}

# Languages that use characters instead of spaces (no word tokenisation)
_CJK_LANGS = {"ja", "zh", "zh-hans", "zh-hant", "zh-cn", "zh-tw"}

_HEADER_RE = re.compile(
    r"^(\d+)"
    r"[.․։۔．]?"
    r"\s*"
    r"([^\d].*)$",
    re.IGNORECASE,
)

# Any docx "Heading N" style (N=1..9) marks a section boundary. Previously
# only "Heading 2" was recognised, which silently broke localisation packs
# where the translator marked the section name (e.g. "BANNER") with
# Heading 3 — that paragraph then leaked into the content text and every
# OCR run got `MANUAL` because the reference began with "BANNER".
_HEADING_STYLE_RE = re.compile(r"^\s*heading\s*\d+\b", re.IGNORECASE)

_NAME_LEADING_STRIP_RE = re.compile(r"^[\s.․։۔．]+")

# Minimum number of explicit numbered headers a TXT must have before it
# is treated as structured. Banner/marketing copy frequently contains a
# single numeric line ("10 JETON", "5 TICKETS") that looks header-shaped
# but is part of the body — sectioning on a single such line splits the
# banner and breaks downstream comparison.
_MIN_TXT_HEADERS_FOR_SECTIONING = 2


@dataclass
class Section:
    number: Optional[int]
    name: str
    content_text: str
    raw_header: str = ""


@dataclass
class ScoredCandidate:
    section: Section
    score: float
    strict_equal: bool
    soft_equal: bool
    has_placeholder_flag: bool
    warnings: List[str] = field(default_factory=list)


@dataclass
class SelectionResult:
    best: Optional[ScoredCandidate]
    all_candidates: List[ScoredCandidate]
    manual_required: bool
    status: str          # "PASS" | "MANUAL"
    delta: float         # top1.score - top2.score
    reason: str
    reference_confidence: float = 0.0
    score_top1: Optional[float] = None
    score_top2: Optional[float] = None
    confidence_margin: float = 0.0


_SCORE_MIN = -0.8
_SCORE_MAX = 6.0
_MARGIN_MAX = _SCORE_MAX - _SCORE_MIN


def _clamp01(value: float) -> float:
    return max(0.0, min(1.0, value))


def _reference_confidence(top1: Optional[ScoredCandidate], top2: Optional[ScoredCandidate]):
    if top1 is None:
        return 0.0, None, None, 0.0

    s1 = float(top1.score)
    s2 = float(top2.score) if top2 is not None else None

    s1n = _clamp01((s1 - _SCORE_MIN) / (_SCORE_MAX - _SCORE_MIN))

    top1_content = normalize_soft(top1.section.content_text)
    top1_valid = bool(top1_content) and (not top1.has_placeholder_flag)

    if s2 is not None:
        margin = s1 - s2
    else:
        margin = _MARGIN_MAX if top1_valid else 0.0

    mn = _clamp01(margin / _MARGIN_MAX)
    conf = _clamp01(0.7 * s1n + 0.3 * mn)
    return conf, s1, s2, margin


# ─── DOCX parsing ────────────────────────────────────────────

def _parse_header(line: str) -> Optional[tuple]:
    stripped = line.strip()
    m = _HEADER_RE.match(stripped)
    if m:
        name = _NAME_LEADING_STRIP_RE.sub("", m.group(2)).strip()
        return int(m.group(1)), name
    return None


def _cell_to_section(cell_text: str) -> Optional[Section]:
    cell_text = cell_text.strip()
    if not cell_text:
        return None

    lines = cell_text.splitlines()
    for i, line in enumerate(lines):
        if not line.strip():
            continue
        parsed = _parse_header(line)
        if parsed:
            num, name = parsed
            content = "\n".join(lines[i + 1:]).strip()
            return Section(number=num, name=name,
                           content_text=content, raw_header=line.strip())
        else:
            return Section(number=None, name="UNKNOWN",
                           content_text=cell_text, raw_header="")
    return None


def _text_from_docx_bytes(docx_bytes: bytes) -> List[Section]:
    from docx import Document  # type: ignore
    doc = Document(io.BytesIO(docx_bytes))
    sections = []

    if doc.tables:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    sec = _cell_to_section(cell.text)
                    if sec:
                        sections.append(sec)
    else:
        sections = _parse_sections_from_paragraphs(doc.paragraphs)
        if not sections or (len(sections) == 1 and sections[0].name == "UNKNOWN"):
            lines = [p.text for p in doc.paragraphs]
            sections = _parse_sections_from_lines(lines)

    return sections


def _is_likely_txt_header_name(name: str) -> bool:
    stripped = name.strip()
    if not stripped or stripped[-1] in "!?.,:;":
        return False

    words = stripped.split()
    if len(words) > 4:
        return False

    letters = [ch for ch in stripped if ch.isalpha()]
    if not letters:
        return False

    upper_ratio = sum(1 for ch in letters if ch.isupper()) / len(letters)
    return upper_ratio >= 0.6


def _text_from_txt_bytes(txt_bytes: bytes) -> List[Section]:
    text = txt_bytes.decode("utf-8", errors="replace")
    lines = text.splitlines()
    parsed_sections = _parse_sections_from_lines(
        lines,
        header_validator=lambda _num, name, _raw: _is_likely_txt_header_name(name),
    )

    explicit_headers = [
        sec for sec in parsed_sections
        if sec.number is not None and sec.raw_header
    ]
    if len(explicit_headers) >= _MIN_TXT_HEADERS_FOR_SECTIONING:
        return parsed_sections

    # Single or no header-shaped line — treat the whole TXT as one block.
    # Banner copy commonly has "10 JETON" / "5 TICKETS" continuations that
    # look like headers but aren't. Sectioning on them has been a source
    # of false MANUALs.
    content = text.strip()
    if not content:
        return []
    return [Section(number=None, name="UNKNOWN", content_text=content, raw_header="")]


def _parse_sections_from_paragraphs(paragraphs) -> List[Section]:
    sections: List[Section] = []
    current_lines: List[str] = []
    current_name = "UNKNOWN"
    current_num: Optional[int] = None
    current_header = ""
    found_heading = False

    def flush():
        content = "\n".join(current_lines).strip()
        if content or current_header:
            sections.append(Section(
                number=current_num, name=current_name,
                content_text=content, raw_header=current_header,
            ))

    for para in paragraphs:
        style_name = para.style.name if para.style else ""
        text = para.text.strip()
        if not text:
            continue

        is_heading = bool(_HEADING_STYLE_RE.match(style_name or ""))
        parsed = _parse_header(text) if not is_heading else None

        if is_heading or parsed:
            found_heading = True
            if current_lines or current_header:
                flush()
                current_lines = []
            if parsed:
                current_num, current_name = parsed
                current_header = text
            else:
                p2 = _parse_header(text)
                if p2:
                    current_num, current_name = p2
                else:
                    current_num = None
                    current_name = text
                current_header = text
        else:
            current_lines.append(text)

    if current_lines or current_header:
        flush()

    return sections if found_heading else []


def _parse_sections_from_lines(lines: List[str], header_validator=None) -> List[Section]:
    sections: List[Section] = []
    current_lines: List[str] = []
    current_num: Optional[int] = None
    current_name = "UNKNOWN"
    current_header = ""
    blank_run = 0

    def flush():
        content = "\n".join(current_lines).strip()
        if content or current_header:
            sections.append(Section(
                number=current_num, name=current_name,
                content_text=content, raw_header=current_header,
            ))

    for raw_line in lines:
        stripped = raw_line.strip()
        if not stripped:
            blank_run += 1
            if blank_run >= 2 and current_lines:
                flush()
                current_lines = []
                current_header = ""
                current_num = None
                current_name = "UNKNOWN"
            else:
                current_lines.append("")
            continue

        blank_run = 0
        parsed = _parse_header(stripped)
        if parsed and (header_validator is None or header_validator(parsed[0], parsed[1], stripped)):
            if current_lines or current_header:
                flush()
                current_lines = []
            current_num, current_name = parsed
            current_header = stripped
        else:
            current_lines.append(stripped)

    if current_lines or current_header:
        flush()

    return sections


def extract_sections(file_bytes: bytes, filename: str) -> List[Section]:
    fname_lower = filename.lower()
    if fname_lower.endswith(".docx"):
        return _text_from_docx_bytes(file_bytes)
    return _text_from_txt_bytes(file_bytes)


# ─── Scoring ────────────────────────────────────────────────────────

def _count_tokens(text: str, lang: str) -> int:
    if lang in _CJK_LANGS:
        return len(re.sub(r"\s+", "", text))
    return len(text.split())


def _name_score(name: str) -> float:
    clean = re.sub(r"^[^\w]+", "", name, flags=re.UNICODE)
    first = clean.split()[0].lower() if clean.split() else ""
    first = first.encode('ascii', errors='ignore').decode('ascii')
    if first in HIGH_PRIORITY_NAMES:
        return 1.0
    if first in PENALTY_NAMES:
        return -0.5
    return 0.0


def _score_section(section: Section, ocr_text: str, lang: str) -> ScoredCandidate:
    warnings: List[str] = []
    score = _name_score(section.name)

    content = section.content_text
    placeholder_flag = has_placeholder(content)
    if placeholder_flag:
        score *= 0.5
        warnings.append("has_placeholder")

    cand_soft   = normalize_soft(content)
    cand_tokens = _count_tokens(cand_soft, lang)

    if lang not in _CJK_LANGS and cand_tokens > 50:
        score -= 0.3
        warnings.append("long_text")

    strict_cmp = compare_lines(ocr_text, content, level="strict")
    soft_cmp   = compare_lines(ocr_text, content, level="soft")
    strict_equal = bool(strict_cmp["pass"])
    soft_equal   = bool(soft_cmp["pass"])

    if strict_equal:
        score += 5.0
    elif soft_equal:
        score += 3.0
    else:
        ocr_soft = normalize_soft(ocr_text)
        ocr_tokens_set  = set(ocr_soft.split())
        cand_tokens_set = set(cand_soft.split())
        if ocr_tokens_set and cand_tokens_set:
            overlap = len(ocr_tokens_set & cand_tokens_set)
            union   = len(ocr_tokens_set | cand_tokens_set)
            score  += (overlap / union) * 2.0 if union else 0

    return ScoredCandidate(
        section=section, score=score,
        strict_equal=strict_equal, soft_equal=soft_equal,
        has_placeholder_flag=placeholder_flag, warnings=warnings,
    )


# ─── Selection ───────────────────────────────────────────────────────

def select_best(
    sections: List[Section],
    ocr_text: str,
    lang: str,
    hint_number: Optional[int] = None,
    hint_name: Optional[str] = None,
) -> SelectionResult:
    if not sections:
        return SelectionResult(best=None, all_candidates=[], manual_required=True,
                               status="MANUAL", delta=0.0, reason="no_sections",
                               reference_confidence=0.0, score_top1=None, score_top2=None, confidence_margin=0.0)

    if len(ocr_text.strip()) < 3:
        return SelectionResult(best=None, all_candidates=[], manual_required=True,
                               status="MANUAL", delta=0.0, reason="ocr_too_short",
                               reference_confidence=0.0, score_top1=None, score_top2=None, confidence_margin=0.0)

    filtered = sections
    if hint_number is not None:
        by_num = [s for s in sections if s.number == hint_number]
        if by_num:
            filtered = by_num
    if hint_name and len(filtered) > 1:
        by_name = [s for s in filtered if hint_name.lower() in s.name.lower()]
        if by_name:
            filtered = by_name

    candidates = [_score_section(s, ocr_text, lang) for s in filtered]
    candidates.sort(key=lambda c: c.score, reverse=True)

    top1  = candidates[0]
    top2  = candidates[1] if len(candidates) > 1 else None
    delta = (top1.score - top2.score) if top2 else 999.0
    conf, s1, s2, margin = _reference_confidence(top1, top2)

    if all(c.has_placeholder_flag for c in candidates):
        if hint_number is not None and len(filtered) == 1:
            pass
        else:
            return SelectionResult(best=top1, all_candidates=candidates, manual_required=True,
                                   status="MANUAL", delta=delta, reason="all_placeholders",
                                   reference_confidence=0.0, score_top1=s1, score_top2=s2, confidence_margin=margin)

    if delta < 0.05 and not top1.strict_equal:
        return SelectionResult(best=top1, all_candidates=candidates, manual_required=True,
                               status="MANUAL", delta=delta, reason="ambiguous_delta",
                               reference_confidence=conf, score_top1=s1, score_top2=s2, confidence_margin=margin)

    if top1.strict_equal:
        return SelectionResult(best=top1, all_candidates=candidates, manual_required=False,
                               status="PASS", delta=delta, reason="strict_equal",
                               reference_confidence=conf, score_top1=s1, score_top2=s2, confidence_margin=margin)
    if top1.soft_equal:
        return SelectionResult(best=top1, all_candidates=candidates, manual_required=True,
                               status="MANUAL", delta=delta, reason="soft_equal_only",
                               reference_confidence=conf, score_top1=s1, score_top2=s2, confidence_margin=margin)

    return SelectionResult(best=top1, all_candidates=candidates, manual_required=False,
                           status="MANUAL", delta=delta, reason="no_match",
                           reference_confidence=conf, score_top1=s1, score_top2=s2, confidence_margin=margin)
