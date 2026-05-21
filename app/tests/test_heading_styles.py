"""
Regression test for the Heading-N section header bug.

Real-world localisation packs (BANNER BNG-30996 archive) mark the
section name (e.g. "BANNER", "バナー") with the `Heading 3` style. The
previous code only recognised `Heading 2`, so the section marker
leaked into the reference content text and every comparison failed.

After the fix, any `Heading N` style is treated as a section boundary
and the marker text becomes `section.name`, not part of `content_text`.
"""
import io

import pytest

from docx import Document  # type: ignore

from app.section_matcher import (
    _parse_sections_from_paragraphs,
    extract_sections,
)


def _make_docx_with_heading(heading_style: str, heading_text: str, body_paragraphs: list[str]) -> bytes:
    """Build a tiny DOCX where the first paragraph uses `heading_style`."""
    doc = Document()
    head = doc.add_paragraph(heading_text)
    head.style = doc.styles[heading_style]
    for line in body_paragraphs:
        doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@pytest.mark.parametrize("style", ["Heading 1", "Heading 2", "Heading 3", "Heading 4"])
def test_any_heading_n_starts_section(style):
    """Heading 1..N all act as section boundaries — only Heading 2 worked previously."""
    docx = _make_docx_with_heading(style, "BANNER", ["", "GIFT COLLECTION", "FOR THE WORLD"])
    sections = extract_sections(docx, "test.docx")
    assert len(sections) == 1
    sec = sections[0]
    assert sec.name == "BANNER"
    # BANNER must NOT leak into the body content
    assert "BANNER" not in sec.content_text
    assert "GIFT COLLECTION" in sec.content_text
    assert "FOR THE WORLD" in sec.content_text


def test_heading3_translated_marker_still_works():
    """Translated banner marker (e.g. Japanese 'バナー') also has Heading 3 style."""
    docx = _make_docx_with_heading("Heading 3", "バナー", ["ギフトコレクション", "世界選手権"])
    sections = extract_sections(docx, "ja.docx")
    assert len(sections) == 1
    assert sections[0].name == "バナー"
    assert "バナー" not in sections[0].content_text
    assert "ギフトコレクション" in sections[0].content_text


def test_no_heading_falls_through_to_unknown():
    """Plain DOCX without any Heading style still works (fallback to line parser)."""
    doc = Document()
    doc.add_paragraph("Just some body text")
    doc.add_paragraph("Another line")
    buf = io.BytesIO()
    doc.save(buf)

    sections = extract_sections(buf.getvalue(), "plain.docx")
    # No header → single UNKNOWN section with all body
    assert len(sections) == 1
    assert sections[0].name == "UNKNOWN"
    assert "Just some body text" in sections[0].content_text
